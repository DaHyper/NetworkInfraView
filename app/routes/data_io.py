import csv
import json
import io
from datetime import datetime, timezone
from flask import Blueprint, render_template, jsonify, request, send_file, flash, redirect, url_for
from app.database import db
from app.models import (Site, ISP, Hardware, Hypervisor, Firewall, VM,
                        App, Storage, Network, ClientDevice, Misc, Connection)

try:
    import yaml
    HAS_YAML = True
except ImportError:
    HAS_YAML = False

try:
    from docx import Document as DocxDocument
    from docx.shared import Inches, Pt
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

bp = Blueprint("data_io", __name__, url_prefix="/io")


def build_export_dict():
    """Build a complete export dictionary from the database."""
    return {
        "export_version": "1.0",
        "exported_at": datetime.now(timezone.utc).isoformat() + "Z",
        "app": "NetworkInfraView",
        "data": {
            "sites":    [s.to_dict() for s in Site.query.all()],
            "isps":     [i.to_dict() for i in ISP.query.all()],
            "hardware": [h.to_dict() for h in Hardware.query.all()],
            "hypervisors": [h.to_dict() for h in Hypervisor.query.all()],
            "firewalls": [f.to_dict() for f in Firewall.query.all()],
            "vms":      [v.to_dict() for v in VM.query.all()],
            "apps":     [a.to_dict() for a in App.query.all()],
            "storage":  [s.to_dict() for s in Storage.query.all()],
            "networks": [n.to_dict() for n in Network.query.all()],
            "clients":  [c.to_dict() for c in ClientDevice.query.all()],
            "misc":     [m.to_dict() for m in Misc.query.all()],
            "connections": [c.to_dict() for c in Connection.query.all()],
        }
    }


def import_from_dict(data, clear_existing=True):
    """Import inventory from a dictionary. Returns (success, message)."""
    if "data" not in data:
        return False, "Invalid format: missing 'data' key."

    d = data["data"]

    if clear_existing:
        # Clear in dependency order
        for model in [Connection, App, VM, Hypervisor, Hardware, Firewall,
                      Storage, ClientDevice, Misc, Network, ISP, Site]:
            db.session.query(model).delete()
        db.session.commit()

    # Build ID maps (old_id → new_obj) for FK remapping
    site_map = {}
    network_map = {}
    hardware_map = {}
    hypervisor_map = {}
    vm_map = {}

    for s in d.get("sites", []):
        old_id = s.pop("id", None)
        obj = Site(**s)
        db.session.add(obj)
        db.session.flush()
        if old_id:
            site_map[old_id] = obj.id

    for i in d.get("isps", []):
        old_id = i.pop("id", None)
        i["site_id"] = site_map.get(i.get("site_id"), i.get("site_id"))
        obj = ISP(**i)
        db.session.add(obj)

    for n in d.get("networks", []):
        old_id = n.pop("id", None)
        n["site_id"] = site_map.get(n.get("site_id"), n.get("site_id"))
        obj = Network(**n)
        db.session.add(obj)
        db.session.flush()
        if old_id:
            network_map[old_id] = obj.id

    for h in d.get("hardware", []):
        old_id = h.pop("id", None)
        h["site_id"] = site_map.get(h.get("site_id"), h.get("site_id"))
        obj = Hardware(**h)
        db.session.add(obj)
        db.session.flush()
        if old_id:
            hardware_map[old_id] = obj.id

    for hy in d.get("hypervisors", []):
        old_id = hy.pop("id", None)
        hy["hardware_id"] = hardware_map.get(hy.get("hardware_id"), hy.get("hardware_id"))
        obj = Hypervisor(**hy)
        db.session.add(obj)
        db.session.flush()
        if old_id:
            hypervisor_map[old_id] = obj.id

    for f in d.get("firewalls", []):
        old_id = f.pop("id", None)
        f["site_id"] = site_map.get(f.get("site_id"), f.get("site_id"))
        obj = Firewall(**f)
        db.session.add(obj)

    for v in d.get("vms", []):
        old_id = v.pop("id", None)
        v["site_id"]       = site_map.get(v.get("site_id"), v.get("site_id"))
        v["network_id"]    = network_map.get(v.get("network_id"), v.get("network_id"))
        v["hypervisor_id"] = hypervisor_map.get(v.get("hypervisor_id"), v.get("hypervisor_id"))
        obj = VM(**v)
        db.session.add(obj)
        db.session.flush()
        if old_id:
            vm_map[old_id] = obj.id

    for a in d.get("apps", []):
        old_id = a.pop("id", None)
        a["vm_id"]       = vm_map.get(a.get("vm_id"), a.get("vm_id"))
        a["hardware_id"] = hardware_map.get(a.get("hardware_id"), a.get("hardware_id"))
        db.session.add(App(**a))

    for s in d.get("storage", []):
        old_id = s.pop("id", None)
        s["site_id"] = site_map.get(s.get("site_id"), s.get("site_id"))
        db.session.add(Storage(**s))

    for c in d.get("clients", []):
        old_id = c.pop("id", None)
        c["site_id"]    = site_map.get(c.get("site_id"), c.get("site_id"))
        c["network_id"] = network_map.get(c.get("network_id"), c.get("network_id"))
        db.session.add(ClientDevice(**c))

    for m in d.get("misc", []):
        old_id = m.pop("id", None)
        m["site_id"] = site_map.get(m.get("site_id"), m.get("site_id"))
        db.session.add(Misc(**m))

    db.session.commit()
    return True, "Import successful."


@bp.route("/")
def index():
    return render_template("io.html")


# ── JSON Export / Import ──────────────────────────────────────────
@bp.route("/export/json")
def export_json():
    data = build_export_dict()
    buf = io.BytesIO(json.dumps(data, indent=2, default=str).encode("utf-8"))
    filename = f"networkinfraview-export-{datetime.now(timezone.utc).strftime('%Y%m%d-%H%M%S')}.json"
    return send_file(buf, mimetype="application/json",
                     as_attachment=True, download_name=filename)


@bp.route("/import/json", methods=["POST"])
def import_json():
    f = request.files.get("file")
    if not f:
        flash("No file uploaded.", "error")
        return redirect(url_for("data_io.index"))
    try:
        data = json.load(f)
        ok, msg = import_from_dict(data)
        flash(msg, "success" if ok else "error")
    except Exception as e:
        flash(f"Import failed: {e}", "error")
    return redirect(url_for("data_io.index"))


# ── YAML Export / Import ──────────────────────────────────────────
@bp.route("/export/yaml")
def export_yaml():
    if not HAS_YAML:
        flash("PyYAML not installed.", "error")
        return redirect(url_for("data_io.index"))
    data = build_export_dict()
    buf = io.BytesIO(yaml.dump(data, default_flow_style=False, allow_unicode=True).encode("utf-8"))
    filename = f"networkinfraview-export-{datetime.now(timezone.utc).strftime('%Y%m%d-%H%M%S')}.yaml"
    return send_file(buf, mimetype="text/yaml",
                     as_attachment=True, download_name=filename)


@bp.route("/import/yaml", methods=["POST"])
def import_yaml():
    if not HAS_YAML:
        flash("PyYAML not installed.", "error")
        return redirect(url_for("data_io.index"))
    f = request.files.get("file")
    if not f:
        flash("No file uploaded.", "error")
        return redirect(url_for("data_io.index"))
    try:
        data = yaml.safe_load(f.read())
        ok, msg = import_from_dict(data)
        flash(msg, "success" if ok else "error")
    except Exception as e:
        flash(f"Import failed: {e}", "error")
    return redirect(url_for("data_io.index"))


# ── DOCX Export ───────────────────────────────────────────────────
@bp.route("/export/docx")
def export_docx():
    if not HAS_DOCX:
        flash("python-docx not installed.", "error")
        return redirect(url_for("data_io.index"))

    doc = DocxDocument()
    doc.add_heading("NetworkInfraView — Infrastructure Report", 0)
    doc.add_paragraph(f"Generated: {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M UTC')}")
    doc.add_paragraph("")

    sections = [
        ("Sites", Site.query.all(), ["Name", "Location", "Address", "Timezone"],
         lambda s: [s.name, s.location or "", s.address or "", s.timezone or ""]),
        ("ISPs", ISP.query.all(), ["Name", "Site", "Type", "ASN", "Public IP Range"],
         lambda i: [i.name, i.site.name, i.type or "", i.asn or "", i.public_ip_range or ""]),
        ("Hardware", Hardware.query.all(), ["Name", "Site", "Type", "Role", "Mgmt IP", "Status"],
         lambda h: [h.name, h.site.name, h.device_type, h.role or "", h.ip_mgmt or "", h.status]),
        ("Firewalls", Firewall.query.all(), ["Name", "Site", "Public IP", "Model", "Status"],
         lambda f: [f.name, f.site.name, f.public_ip or "", f.model or "", f.status]),
        ("Virtual Machines", VM.query.all(), ["Name", "Site", "OS", "IP", "Network", "Status"],
         lambda v: [v.name, v.site.name, v.os or "", v.ip_address or "",
                    (v.network.name if v.network else ""), v.status]),
        ("Apps & Services", App.query.all(), ["Name", "Host", "Port", "Protocol", "Public"],
         lambda a: [a.name, (a.vm.name if a.vm else a.hardware.name if a.hardware else ""),
                    str(a.port or ""), a.protocol or "", "Yes" if a.public_exposed else "No"]),
        ("Storage", Storage.query.all(), ["Name", "Site", "Type", "Capacity TB", "IP"],
         lambda s: [s.name, s.site.name, s.type or "", str(s.capacity_tb or ""), s.ip_address or ""]),
        ("Networks / VLANs", Network.query.all(), ["Name", "Site", "VLAN ID", "Subnet"],
         lambda n: [n.name, n.site.name, str(n.vlan_id or ""), n.subnet or ""]),
        ("Client Devices", ClientDevice.query.all(), ["Name", "Site", "Owner", "Type", "IP"],
         lambda c: [c.name, c.site.name, c.owner or "", c.device_type or "", c.ip_address or ""]),
        ("Misc", Misc.query.all(), ["Name", "Site", "Category", "IP"],
         lambda m: [m.name, m.site.name, m.category or "", m.ip_address or ""]),
    ]

    for title, items, headers, row_fn in sections:
        doc.add_heading(title, level=1)
        if not items:
            doc.add_paragraph("(none)")
            continue
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = "Table Grid"
        hdr_row = table.rows[0]
        for i, h in enumerate(headers):
            hdr_row.cells[i].text = h
            run = hdr_row.cells[i].paragraphs[0].runs[0]
            run.bold = True
        for item in items:
            row = table.add_row()
            for i, val in enumerate(row_fn(item)):
                row.cells[i].text = val
        doc.add_paragraph("")

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    filename = f"networkinfraview-report-{datetime.now(timezone.utc).strftime('%Y%m%d')}.docx"
    return send_file(buf,
                     mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                     as_attachment=True, download_name=filename)


# ── CSV Import ────────────────────────────────────────────────────
# Expected CSV column headers per entity type:
#   hardware: name, site_name, device_type, role, ip_mgmt, cpu_cores, ram_gb, make_model, status, notes
#   vms:      name, site_name, os, ip_address, cpu_cores, ram_gb, role, status, notes
#   networks: name, site_name, vlan_id, subnet, color, description
#   sites:    name, location, address, timezone, notes
#   clients:  name, site_name, device_type, owner, ip_address, mac_address, os, notes

@bp.route("/import/csv", methods=["POST"])
def import_csv():
    f = request.files.get("file")
    entity = request.form.get("entity", "").strip()
    if not f or not entity:
        flash("File and entity type required.", "error")
        return redirect(url_for("data_io.index"))

    try:
        text = f.read().decode("utf-8-sig")
        reader = csv.DictReader(io.StringIO(text))
        rows = list(reader)
        if not rows:
            flash("CSV is empty.", "error")
            return redirect(url_for("data_io.index"))

        # Build site lookup by name
        site_lookup = {s.name.lower(): s.id for s in Site.query.all()}

        def get_site(name):
            return site_lookup.get((name or "").strip().lower())

        count = 0
        if entity == "hardware":
            for r in rows:
                sid = get_site(r.get("site_name", ""))
                if not sid:
                    continue
                db.session.add(Hardware(
                    site_id=sid,
                    name=r.get("name", "").strip(),
                    device_type=r.get("device_type", "server").strip() or "server",
                    role=r.get("role", "").strip(),
                    ip_mgmt=r.get("ip_mgmt", "").strip() or None,
                    cpu_cores=int(r.get("cpu_cores") or 0) or None,
                    ram_gb=int(r.get("ram_gb") or 0) or None,
                    make_model=r.get("make_model", "").strip(),
                    status=r.get("status", "active").strip() or "active",
                    notes=r.get("notes", "").strip(),
                ))
                count += 1

        elif entity == "vms":
            hyp_lookup = {h.hardware.name.lower(): h.id
                          for h in Hypervisor.query.all() if h.hardware}
            net_lookup = {n.name.lower(): n.id for n in Network.query.all()}
            for r in rows:
                sid = get_site(r.get("site_name", ""))
                if not sid:
                    continue
                db.session.add(VM(
                    site_id=sid,
                    hypervisor_id=hyp_lookup.get(r.get("hypervisor_name", "").lower()),
                    network_id=net_lookup.get(r.get("network_name", "").lower()),
                    name=r.get("name", "").strip(),
                    os=r.get("os", "").strip(),
                    ip_address=r.get("ip_address", "").strip() or None,
                    cpu_cores=int(r.get("cpu_cores") or 0) or None,
                    ram_gb=int(r.get("ram_gb") or 0) or None,
                    role=r.get("role", "").strip(),
                    status=r.get("status", "active").strip() or "active",
                    notes=r.get("notes", "").strip(),
                ))
                count += 1

        elif entity == "networks":
            for r in rows:
                sid = get_site(r.get("site_name", ""))
                if not sid:
                    continue
                db.session.add(Network(
                    site_id=sid,
                    name=r.get("name", "").strip(),
                    vlan_id=int(r.get("vlan_id") or 0) or None,
                    subnet=r.get("subnet", "").strip() or None,
                    color=r.get("color", "#6366f1").strip() or "#6366f1",
                    description=r.get("description", "").strip(),
                ))
                count += 1

        elif entity == "sites":
            for r in rows:
                db.session.add(Site(
                    name=r.get("name", "").strip(),
                    location=r.get("location", "").strip(),
                    address=r.get("address", "").strip(),
                    timezone=r.get("timezone", "").strip(),
                    notes=r.get("notes", "").strip(),
                ))
                count += 1

        elif entity == "clients":
            net_lookup = {n.name.lower(): n.id for n in Network.query.all()}
            for r in rows:
                sid = get_site(r.get("site_name", ""))
                if not sid:
                    continue
                db.session.add(ClientDevice(
                    site_id=sid,
                    network_id=net_lookup.get(r.get("network_name", "").lower()),
                    name=r.get("name", "").strip(),
                    owner=r.get("owner", "").strip(),
                    device_type=r.get("device_type", "laptop").strip() or "laptop",
                    ip_address=r.get("ip_address", "").strip() or None,
                    mac_address=r.get("mac_address", "").strip() or None,
                    os=r.get("os", "").strip(),
                    notes=r.get("notes", "").strip(),
                ))
                count += 1

        else:
            flash(f"Unknown entity type: {entity}", "error")
            return redirect(url_for("data_io.index"))

        db.session.commit()
        flash(f"Imported {count} {entity} from CSV.", "success")
    except Exception as e:
        db.session.rollback()
        flash(f"CSV import failed: {e}", "error")

    return redirect(url_for("data_io.index"))


@bp.route("/export/csv/<entity>")
def export_csv(entity):
    """Export a single entity type as CSV."""
    headers_map = {
        "hardware": (["name", "site", "device_type", "role", "ip_mgmt", "cpu_cores", "ram_gb", "make_model", "status", "notes"],
                     Hardware.query.all(),
                     lambda h: [h.name, h.site.name, h.device_type, h.role or "", h.ip_mgmt or "",
                                str(h.cpu_cores or ""), str(h.ram_gb or ""), h.make_model or "", h.status, h.notes or ""]),
        "vms":      (["name", "site", "os", "ip_address", "cpu_cores", "ram_gb", "role", "status", "notes"],
                     VM.query.all(),
                     lambda v: [v.name, v.site.name, v.os or "", v.ip_address or "",
                                str(v.cpu_cores or ""), str(v.ram_gb or ""), v.role or "", v.status, v.notes or ""]),
        "networks": (["name", "site", "vlan_id", "subnet", "color", "description"],
                     Network.query.all(),
                     lambda n: [n.name, n.site.name, str(n.vlan_id or ""), n.subnet or "", n.color or "", n.description or ""]),
        "sites":    (["name", "location", "address", "timezone", "notes"],
                     Site.query.all(),
                     lambda s: [s.name, s.location or "", s.address or "", s.timezone or "", s.notes or ""]),
        "clients":  (["name", "site", "device_type", "owner", "ip_address", "mac_address", "os", "notes"],
                     ClientDevice.query.all(),
                     lambda c: [c.name, c.site.name, c.device_type or "", c.owner or "",
                                c.ip_address or "", c.mac_address or "", c.os or "", c.notes or ""]),
    }
    if entity not in headers_map:
        flash("Unknown entity.", "error")
        return redirect(url_for("data_io.index"))

    headers, items, row_fn = headers_map[entity]
    buf = io.StringIO()
    w = csv.writer(buf)
    w.writerow(headers)
    for item in items:
        w.writerow(row_fn(item))

    out = io.BytesIO(buf.getvalue().encode("utf-8"))
    filename = f"networkinfraview-{entity}-{datetime.now(timezone.utc).strftime('%Y%m%d')}.csv"
    return send_file(out, mimetype="text/csv", as_attachment=True, download_name=filename)


# ── API export for JS ─────────────────────────────────────────────
@bp.route("/api/export/json")
def api_export_json():
    return jsonify(build_export_dict())
